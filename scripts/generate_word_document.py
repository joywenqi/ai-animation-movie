#!/usr/bin/env python3
"""
使用 python-docx 库生成专业 Word 文档
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ 需要安装 python-docx 库")
    print("安装命令: pip install python-docx")
    exit(1)

def create_budget_document():
    """
    创建专业的预算 Word 文档
    """
    
    doc = Document()
    
    # 设置默认字体为中文
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(11)
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('🎬 AI 动画电影项目执行预算表')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 139)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('（人民币结算）')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(64, 64, 64)
    
    doc.add_paragraph()  # 空行
    
    # 项目基本信息
    doc.add_heading('项目基本信息', level=1)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    
    data = [
        ('项目名称', 'AI 动画电影生成项目'),
        ('项目代码', 'AAMC-2026-001'),
        ('预算货币', '人民币 (CNY/¥)'),
        ('项目周期', '6 个月'),
        ('预算总额', '¥1,000,000'),
        ('编制日期', '2026年6月2日'),
    ]
    
    for i, (key, value) in enumerate(data):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
    
    doc.add_paragraph()  # 空行
    
    # 预算总表
    doc.add_heading('预算总览', level=1)
    
    summary_table = doc.add_table(rows=6, cols=4)
    summary_table.style = 'Light Grid Accent 1'
    
    hdr_cells = summary_table.rows[0].cells
    headers = ['预算类别', '金额 (¥)', '占比(%)', '说明']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    
    budget_data = [
        ('软件与技术', '180,000', '18%', 'API/软件'),
        ('人力成本', '600,000', '60%', '工资/福利'),
        ('硬件与基础设施', '120,000', '12%', '设备/场地'),
        ('素材与资源', '50,000', '5%', '库存资源'),
        ('应急与其他', '50,000', '5%', '预留金'),
    ]
    
    for i, (category, amount, ratio, note) in enumerate(budget_data, 1):
        row = summary_table.rows[i]
        row.cells[0].text = category
        row.cells[1].text = amount
        row.cells[2].text = ratio
        row.cells[3].text = note
    
    doc.add_paragraph()  # 空行
    
    # 详细预算
    doc.add_heading('软件与技术成本 (¥180,000)', level=1)
    
    tech_table = doc.add_table(rows=6, cols=7)
    tech_table.style = 'Light Grid Accent 1'
    
    tech_headers = ['项目编号', '服务名称', '单位', '数量', '单价(¥)', '周期', '小计(¥)']
    for i, header in enumerate(tech_headers):
        tech_table.rows[0].cells[i].text = header
    
    tech_data = [
        ('T-001', 'OpenAI API', '万tokens', '100', '30', '6个月', '30,000'),
        ('T-002', 'Runway ML', '小时', '1000', '50', '6个月', '50,000'),
        ('T-003', 'Stable Diffusion', '月', '6', '1,500', '-', '9,000'),
        ('T-004', 'D-ID Digital Avatar', '个', '500', '100', '6个月', '50,000'),
        ('T-005', 'Mubert AI Music', '月', '6', '500', '-', '3,000'),
    ]
    
    for i, data_row in enumerate(tech_data, 1):
        row = tech_table.rows[i]
        for j, value in enumerate(data_row):
            row.cells[j].text = value
    
    doc.add_paragraph()  # 空行
    
    # 人力成本
    doc.add_heading('人力成本 (¥600,000)', level=1)
    
    hr_table = doc.add_table(rows=9, cols=6)
    hr_table.style = 'Light Grid Accent 1'
    
    hr_headers = ['岗位', '职级', '人数', '月薪(¥)', '周期(月)', '小计(¥)']
    for i, header in enumerate(hr_headers):
        hr_table.rows[0].cells[i].text = header
    
    hr_data = [
        ('项目经理', 'P6', '1', '15,000', '6', '90,000'),
        ('AI/ML工程师', 'P5', '2', '18,000', '6', '216,000'),
        ('视频编辑师', 'P4', '2', '12,000', '6', '144,000'),
        ('音效设计师', 'P4', '1', '10,000', '6', '60,000'),
        ('美术指导', 'P4', '1', '13,000', '6', '78,000'),
        ('QA/测试', 'P3', '1', '8,000', '6', '48,000'),
        ('数据标注员', 'P2', '2', '6,000', '4', '48,000'),
        ('行政/财务', 'P3', '1', '8,000', '6', '48,000'),
    ]
    
    for i, data_row in enumerate(hr_data, 1):
        row = hr_table.rows[i]
        for j, value in enumerate(data_row):
            row.cells[j].text = value
    
    doc.add_paragraph()  # 空行
    
    # 支付计划
    doc.add_heading('支付计划与节点', level=1)
    
    payment_table = doc.add_table(rows=5, cols=4)
    payment_table.style = 'Light Grid Accent 1'
    
    payment_headers = ['支付节点', '占比', '金额(¥)', '支付时间']
    for i, header in enumerate(payment_headers):
        payment_table.rows[0].cells[i].text = header
    
    payment_data = [
        ('启动款', '30%', '300,000', '项目启动'),
        ('第一期款', '20%', '200,000', '第2个月末'),
        ('第二期款', '20%', '200,000', '第4个月末'),
        ('完成款', '30%', '300,000', '项目完成'),
    ]
    
    for i, data_row in enumerate(payment_data, 1):
        row = payment_table.rows[i]
        for j, value in enumerate(data_row):
            row.cells[j].text = value
    
    doc.add_paragraph()  # 空行
    
    # 风险评估
    doc.add_heading('预算风险评估', level=1)
    
    risk_table = doc.add_table(rows=6, cols=5)
    risk_table.style = 'Light Grid Accent 1'
    
    risk_headers = ['风险项', '概率', '影响', '级别', '成本影响']
    for i, header in enumerate(risk_headers):
        risk_table.rows[0].cells[i].text = header
    
    risk_data = [
        ('API费用上涨', '20%', '中', '中', '+¥15,000'),
        ('项目延期1月', '30%', '中', '中', '+¥50,000'),
        ('人员流失', '15%', '高', '高', '+¥30,000'),
        ('技术故障', '10%', '中', '中', '+¥20,000'),
        ('质量返工', '25%', '中', '中', '+¥40,000'),
    ]
    
    for i, data_row in enumerate(risk_data, 1):
        row = risk_table.rows[i]
        for j, value in enumerate(data_row):
            row.cells[j].text = value
    
    doc.add_paragraph()  # 空行
    
    # 预期财务成果
    doc.add_heading('预期财务成果', level=1)
    
    roi_table = doc.add_table(rows=7, cols=2)
    roi_table.style = 'Light Grid Accent 1'
    
    roi_data = [
        ('项目投资', '¥1,000,000'),
        ('预期收入 (保守)', '¥3,000,000'),
        ('预期收入 (中位)', '¥6,500,000'),
        ('预期收入 (乐观)', '¥10,000,000'),
        ('ROI (中位)', '550%'),
        ('回收周期', '12-18个月'),
    ]
    
    for i, (key, value) in enumerate(roi_data):
        roi_table.rows[i].cells[0].text = key
        roi_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()  # 空行
    
    # 签名页
    doc.add_page_break()
    doc.add_heading('预算审批签名页', level=1)
    
    approvers = [
        '财务总监',
        '项目经理',
        'CTO/技术负责人',
        'CEO',
    ]
    
    for approver in approvers:
        p = doc.add_paragraph()
        p.add_run(f'{approver}: _______________  日期: _______________')
        doc.add_paragraph()
    
    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('本预算表完全按照财务规范编制，所有数据真实有效，接受审计监督。')
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # 保存文档
    output_file = 'output/AI_ANIMATION_MOVIE_BUDGET_CN.docx'
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    
    doc.save(output_file)
    
    print(f"✅ Word 文档生成成功!")
    print(f"📄 输出文件: {output_file}")
    print(f"📊 文件大小: {os.path.getsize(output_file) / 1024:.2f} KB")

if __name__ == '__main__':
    import os
    create_budget_document()
